from docx import Document
from builders.common import *

def build(data):
    doc = Document()
    set_page_margins(doc)
    
    add_header_table(doc, data)
    add_title(doc, f"Bổ sung và thay thế cán bộ đoàn {data['loai_tu_van']}")
    add_tgd(doc)
    add_can_cu(doc, data)
    add_quyet_dinh_label(doc)
    
    # Điều 1
    p1 = build_dieu1_header(doc, "Bổ sung và thay thế cán bộ", data['loai_tu_van'], data)
    r_tail = p1.add_run(" cụ thể như sau:")
    fmt_body(r_tail)
    
    for i, m in enumerate(data["members"], 1):
        p = doc.add_paragraph()
        para_spacing_body(p)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        
        r_m = p.add_run(f"{i}. {m['trinh_do']}. {m['ho_ten']}\t–\t{m['chuc_vu']}.")
        fmt_body(r_m)
        
        if m.get("la_thay_the") and m.get("ho_ten_cu"):
            p2 = doc.add_paragraph()
            para_spacing_body(p2)
            p2.paragraph_format.left_indent = Cm(2.5)
            r_replace = p2.add_run(f"(thay thế {m['trinh_do_cu']}. {m['ho_ten_cu']})")
            fmt_body(r_replace)
    
    add_dieu_2_3_4(doc, data["ten_tt"], data["loai_tu_van"])
    add_footer_table(doc)
    return doc
