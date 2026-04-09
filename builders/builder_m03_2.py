from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from builders.common import (
    set_page_margins, add_header_table, add_title, add_tgd,
    add_can_cu, add_quyet_dinh_label, build_dieu1_header,
    add_dieu_2_3_4, add_footer_table,
    set_cell_vertical_align_center,
    para_spacing_body, para_spacing_table, fmt_body, FONT_SIZE_BODY
)

COL_WIDTHS_M03_2 = [Mm(12), Mm(60), Mm(45), Mm(43)]

def build(data):
    doc = Document()
    set_page_margins(doc)

    add_header_table(doc, data)
    add_title(doc, f"Bổ sung và thay thế cán bộ đoàn {data['loai_tu_van']}")
    add_tgd(doc)
    add_can_cu(doc, data)
    add_quyet_dinh_label(doc)

    p1 = build_dieu1_header(doc, "Bổ sung và thay thế cán bộ", data["loai_tu_van"], data)
    fmt_body(p1.add_run(" gồm các cán bộ theo danh sách đính kèm."))

    add_dieu_2_3_4(doc, data["ten_tt"], data["loai_tu_van"])
    add_footer_table(doc)

    doc.add_page_break()
    _add_danh_sach_page_m03_2(doc, data)
    return doc


def _add_danh_sach_page_m03_2(doc, data):
    p_ds = doc.add_paragraph()
    p_ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing_body(p_ds, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_ds.add_run("DANH SÁCH BỔ SUNG VÀ THAY THẾ CÁN BỘ ĐOÀN TƯ VẤN"), bold=True)

    p_kt = doc.add_paragraph()
    para_spacing_body(p_kt, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_kt.add_run(
        f"Kèm theo Quyết định số: {data['so_qd_display']} ngày {data['ngay_qd_display']}"
    ))

    doc.add_paragraph()   # dòng trống

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, w in enumerate(COL_WIDTHS_M03_2):
        table.columns[i].width = w
        table.rows[0].cells[i].width = w

    headers = ["STT", "Họ và tên", "Vị trí công việc", "Cán bộ bị thay thế"]
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt_body(p.add_run(h), bold=True)
        set_cell_vertical_align_center(table.rows[0].cells[i])

    for idx, m in enumerate(data["members"], 1):
        row = table.add_row()
        for i, w in enumerate(COL_WIDTHS_M03_2):
            row.cells[i].width = w

        # Cột STT
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p0, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt_body(p0.add_run(str(idx)))
        set_cell_vertical_align_center(row.cells[0])

        # Cột Họ tên cán bộ mới
        p1 = row.cells[1].paragraphs[0]
        para_spacing_table(p1)
        fmt_body(p1.add_run(f"{m['trinh_do']}. {m['ho_ten']}"))
        set_cell_vertical_align_center(row.cells[1])

        # Cột Vị trí
        p2 = row.cells[2].paragraphs[0]
        para_spacing_table(p2)
        fmt_body(p2.add_run(m["chuc_vu"]))
        set_cell_vertical_align_center(row.cells[2])

        # Cột Cán bộ bị thay thế – chỉ điền nếu là thay thế, để trống nếu chỉ bổ sung
        p3 = row.cells[3].paragraphs[0]
        para_spacing_table(p3)
        if m.get("la_thay_the") and m.get("ho_ten_cu", "").strip():
            fmt_body(p3.add_run(f"{m['trinh_do_cu']}. {m['ho_ten_cu']}"))
        set_cell_vertical_align_center(row.cells[3])
