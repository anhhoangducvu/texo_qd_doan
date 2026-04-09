from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
FONT_SIZE_HEADER = Pt(12)   # vùng header (quốc hiệu) và footer (ký tên)
FONT_SIZE_BODY = Pt(13)     # vùng thân văn bản
FONT_SIZE_TITLE = Pt(14)    # Tiêu đề QUYẾT ĐỊNH

def remove_cell_borders(cell):
    """Xóa viền ô trong bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph_bottom_border(paragraph, space_pt=3):
    """
    Kẻ đường kẻ dưới đoạn văn bằng OOXML pBdr/bottom.
    space_pt: khoảng cách từ chân chữ đến đường kẻ.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single') # đường kẻ đơn
    bottom.set(qn('w:sz'), '6')      # độ dày 0.75pt
    bottom.set(qn('w:space'), str(space_pt * 20)) # twips
    bottom.set(qn('w:color'), 'auto')

    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_vertical_align_center(cell):
    """Căn chữ giữa ô theo chiều cao (vertical center)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def set_page_margins(doc):
    """Thiết lập khổ giấy A4 và căn lề chuẩn văn bản hành chính VN."""
    for section in doc.sections:
        # Khổ giấy A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Căn lề
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def fmt_header(run, bold=False):
    """Font 12pt, dùng cho bảng header quốc hiệu và bảng footer ký tên."""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_HEADER
    run.font.bold = bold
    run.font.italic = False

def fmt_body(run, bold=False, italic=False, size=None):
    """Font 13pt (mặc định), dùng cho toàn bộ thân văn bản."""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = size or FONT_SIZE_BODY
    run.font.bold = bold
    run.font.italic = italic

def para_spacing_header(paragraph, align=None):
    """Giãn dòng Exactly 15pt, trên/dưới 0pt – dùng cho header & footer bảng."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(15)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align

def para_spacing_body(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    Giãn dòng Exactly 17pt, trước 6pt, sau 0pt.
    Mặc định căn JUSTIFY.
    """
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(17)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    paragraph.alignment = align

def para_spacing_table(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    Spacing dành riêng cho nội dung bên trong bảng đính kèm.
    Before 3pt / After 3pt / Line spacing Exactly 15pt.
    """
    pf = paragraph.paragraph_format
    pf.space_before      = Pt(3)
    pf.space_after       = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(15)
    paragraph.alignment  = align

def add_header_table(doc, data):
    """Bảng header 2 cột không viền: logo trái, quốc hiệu phải"""
    table = doc.add_table(rows=1, cols=2)
    # Cố định chiều rộng cột
    col_left = table.columns[0]
    col_right = table.columns[1]
    col_left.width = Mm(60)
    col_right.width = Mm(100)
    
    for row in table.rows:
        row.cells[0].width = Mm(60)
        row.cells[1].width = Mm(100)
        for cell in row.cells:
            remove_cell_borders(cell)
    
    # --- Cột trái ---
    c_left = table.cell(0, 0)
    p1 = c_left.paragraphs[0]
    para_spacing_header(p1, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p1.add_run("CÔNG TY CỔ PHẦN TEXO"), bold=True)
    
    p2 = c_left.add_paragraph()
    para_spacing_header(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p2.add_run("TƯ VẤN VÀ ĐẦU TƯ"), bold=True)
    
    p_blank_l = c_left.add_paragraph() # Dòng trống đối xứng
    para_spacing_header(p_blank_l)
    
    p3 = c_left.add_paragraph()
    para_spacing_header(p3, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p3.add_run(f"Số: {data['so_qd_display']}"))
    
    # --- Cột phải ---
    c_right = table.cell(0, 1)
    p4 = c_right.paragraphs[0]
    para_spacing_header(p4, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p4.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM"), bold=True)
    
    p5 = c_right.add_paragraph()
    para_spacing_header(p5, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p5.add_run("Độc lập - Tự do - Hạnh phúc"), bold=True)
    
    p6 = c_right.add_paragraph()  # dòng trống
    para_spacing_header(p6)
    
    p7 = c_right.add_paragraph()
    para_spacing_header(p7, align=WD_ALIGN_PARAGRAPH.RIGHT) # Căn PHẢI dòng ngày
    fmt_header(p7.add_run(f"Hà Nội, {data['ngay_qd_display']}"))

def add_title(doc, tieu_de_ve_viec):
    """QUYẾT ĐỊNH + Về việc..."""
    p_empty = doc.add_paragraph()
    para_spacing_body(p_empty, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_qd = doc.add_paragraph()
    para_spacing_body(p_qd, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_qd.add_run("QUYẾT ĐỊNH"), bold=True, size=FONT_SIZE_TITLE)
    
    p_vv = doc.add_paragraph()
    para_spacing_body(p_vv, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_vv.add_run(f"Về việc: {tieu_de_ve_viec}"), bold=True)

def add_tgd(doc):
    p = doc.add_paragraph()
    para_spacing_body(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p.add_run("TỔNG GIÁM ĐỐC CÔNG TY"), bold=True)

def add_can_cu(doc, data):
    """Các dòng căn cứ"""
    can_cu_list = [
        "Căn cứ Điều lệ của Công ty Cổ phần TEXO Tư vấn và Đầu tư;",
        f"Căn cứ hợp đồng kinh tế số: {data['so_hd_cdt']} ({data['so_hd_texo']}) ngày {data['ngay_hd']} giữa {data['ten_cdt']} với Công ty Cổ phần TEXO Tư vấn và Đầu Tư, về việc {data['noi_dung_hd']};",
        "Căn cứ vào năng lực cán bộ và nhu cầu công tác;",
        f"Căn cứ vào đề nghị của Giám đốc {data['ten_tt']};",
        "Theo đề nghị của Phòng Kỹ thuật.",
    ]
    for cc in can_cu_list:
        p = doc.add_paragraph()
        para_spacing_body(p) # Mặc định JUSTIFY
        r = p.add_run(f"- {cc}")
        fmt_body(r, italic=True)

def add_quyet_dinh_label(doc):
    p = doc.add_paragraph()
    para_spacing_body(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p.add_run("QUYẾT ĐỊNH"), bold=True)

def build_dieu1_header(doc, hanh_dong, loai_tu_van, data):
    """Hàm helper xây dựng phần đầu Điều 1."""
    p1 = doc.add_paragraph()
    para_spacing_body(p1) # Mặc định JUSTIFY
    p1.paragraph_format.first_line_indent = Cm(1.0) # Thụt đầu dòng 1cm

    r_label = p1.add_run("Điều 1:")
    fmt_body(r_label, bold=True)

    r_content = p1.add_run(
        f" {hanh_dong} đoàn {loai_tu_van} thực hiện hợp đồng số:"
        f" {data['so_hd_cdt']} ({data['so_hd_texo']})"
    )
    fmt_body(r_content)
    return p1

def format_member_line(doc, idx, member):
    """Sinh dòng danh sách thành viên."""
    p = doc.add_paragraph()
    para_spacing_body(p) # Mặc định JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5) # Hanging indent cho số thứ tự
    
    r = p.add_run(f"{idx}. {member['trinh_do']}. {member['ho_ten']}\t–\t{member['chuc_vu']}.")
    fmt_body(r)

def build_lead_title(loai_tu_van: str) -> str:
    """
    Trả về chức danh trưởng đoàn theo loại dịch vụ tư vấn.
    """
    s = loai_tu_van.lower()
    if any(k in s for k in ["giám sát", "tvgs"]):
        return "Giám sát trưởng"
    if any(k in s for k in ["quản lý dự án", "qlda", "tvqlda", "tv qlda"]):
        return "Giám đốc Quản lý dự án"
    if any(k in s for k in ["thẩm tra", "tvtt", "tv thẩm tra"]):
        return "Chủ trì"
    if any(k in s for k in ["kiểm định", "tvkđ"]):
        return "Chủ trì"
    return "Trưởng đoàn"

def build_dieu2_content(loai_tu_van: str) -> str:
    """
    Trả về nội dung thân Điều 2 phù hợp with loại dịch vụ tư vấn.
    """
    s = loai_tu_van.lower()
    lead = build_lead_title(loai_tu_van)

    if any(k in s for k in ["giám sát", "tvgs"]):
        return (
            f" {lead} và các thành viên trong đoàn chịu trách nhiệm công việc"
            " giám sát theo đúng quy định của Pháp luật hiện hành về tư vấn giám sát"
            " xây dựng thi công công trình, các điều khoản ghi trong hợp đồng kinh tế"
            " và đề cương tư vấn giám sát đã ký với khách hàng. Thường xuyên có mặt tại"
            " công trình theo đúng tiến độ thoả thuận với khách hàng và báo cáo đầy đủ"
            " về Công ty theo quy định."
        )

    if any(k in s for k in ["quản lý dự án", "qlda", "tvqlda", "tv qlda"]):
        return (
            f" {lead} và các thành viên trong đoàn chịu trách nhiệm công việc quản lý"
            " dự án theo đúng quy định của Pháp luật hiện hành về tư vấn quản lý dự án"
            " xây dựng, các điều khoản ghi trong hợp đồng kinh tế và đề cương tư vấn"
            " quản lý dự án đã ký với khách hàng. Thường xuyên thực hiện đúng nhiệm vụ"
            " được phân công và báo cáo đầy đủ về Công ty theo quy định."
        )

    if any(k in s for k in ["thẩm tra", "tvtt", "tv thẩm tra"]):
        return (
            f" {lead} và các thành viên trong đoàn chịu trách nhiệm công việc thẩm tra"
            " theo đúng quy định của Pháp luật hiện hành về tư vấn thẩm tra thiết kế"
            " xây dựng, các điều khoản ghi trong hợp đồng kinh tế và đề cương tư vấn"
            " thẩm tra đã ký với khách hàng. Hoàn thành đúng tiến độ thoả thuận với"
            " khách hàng và báo cáo đầy đủ về Công ty theo quy định."
        )

    if any(k in s for k in ["kiểm định", "tvkđ"]):
        return (
            f" {lead} và các thành viên trong đoàn chịu trách nhiệm công việc kiểm định"
            " chất lượng công trình theo đúng quy định của Pháp luật hiện hành, các điều"
            " khoản ghi trong hợp đồng kinh tế và đề cương đã ký với khách hàng. Hoàn"
            " thành đúng tiến độ thoả thuận với khách hàng và báo cáo đầy đủ về Công ty"
            " theo quy định."
        )

    # Fallback
    return (
        f" {lead} và các thành viên chịu trách nhiệm thực hiện công việc"
        f" {loai_tu_van.lower()} theo đúng quy định của Pháp luật hiện hành, các điều"
        f" khoản ghi trong hợp đồng kinh tế và đề cương đã ký với khách hàng. Hoàn"
        f" thành đúng tiến độ thoả thuận với khách hàng và báo cáo đầy đủ về Công ty"
        f" theo quy định."
    )

def add_dieu_2_3_4(doc, ten_tt, loai_tu_van):
    """Điều 2, 3, 4"""
    lead = build_lead_title(loai_tu_van)
    
    # Điều 2
    p2 = doc.add_paragraph()
    para_spacing_body(p2) # Mặc định JUSTIFY
    p2.paragraph_format.first_line_indent = Cm(1.0)
    fmt_body(p2.add_run("Điều 2:"), bold=True)
    fmt_body(p2.add_run(build_dieu2_content(loai_tu_van)))
    
    # Điều 3
    p3 = doc.add_paragraph()
    para_spacing_body(p3) # Mặc định JUSTIFY
    p3.paragraph_format.first_line_indent = Cm(1.0)
    fmt_body(p3.add_run("Điều 3:"), bold=True)
    fmt_body(p3.add_run(" Quyết định này có hiệu lực từ ngày ký."))
    
    # Điều 4
    p4 = doc.add_paragraph()
    para_spacing_body(p4) # Mặc định JUSTIFY
    p4.paragraph_format.first_line_indent = Cm(1.0)
    fmt_body(p4.add_run("Điều 4:"), bold=True)
    fmt_body(p4.add_run(
        f" Giám đốc {ten_tt}, {lead} và các thành viên,"
        " Trưởng các phòng quản lý chức năng và các cá nhân liên quan"
        " chịu trách nhiệm thi hành Quyết định này./."
    ))
    
    # Dòng trống tạo khoảng cách trước bảng ký tên
    p_blank = doc.add_paragraph()
    para_spacing_body(p_blank)

def add_footer_table(doc):
    """Bảng footer: Nơi nhận | Chữ ký TGĐ"""
    table = doc.add_table(rows=1, cols=2)
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)
    
    # --- Cột trái: Nơi nhận ---
    c_left = table.cell(0, 0)
    p1 = c_left.paragraphs[0]
    para_spacing_header(p1)
    fmt_header(p1.add_run("Nơi nhận:"), bold=True)
    
    for line in ["- Như điều 4;", "- Chủ đầu tư;", "- Lưu: VT."]:
        p = c_left.add_paragraph()
        para_spacing_header(p)
        fmt_header(p.add_run(line))
        
    # --- Cột phải: Chữ ký ---
    c_right = table.cell(0, 1)
    p2 = c_right.paragraphs[0]
    para_spacing_header(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p2.add_run("CÔNG TY CỔ PHẦN TEXO"), bold=True)
    
    p3 = c_right.add_paragraph()
    para_spacing_header(p3, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_header(p3.add_run("TƯ VẤN VÀ ĐẦU TƯ"), bold=True)
    
    for _ in range(5): # Đã tăng lên 5 dòng trống
        p_blank = c_right.add_paragraph()
        para_spacing_header(p_blank)
        
    # Đã bỏ dòng "TỔNG GIÁM ĐỐC"
