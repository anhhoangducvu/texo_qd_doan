from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from builders.common import (
    set_page_margins, add_header_table, add_title, add_tgd,
    add_can_cu, add_quyet_dinh_label, build_dieu1_header,
    add_dieu_2_3_4, add_footer_table,
    set_cell_vertical_align_center,
    para_spacing_body, para_spacing_table, fmt_body, FONT_SIZE_BODY
)

COL_WIDTHS = [Mm(15), Mm(85), Mm(60)]

def build(data):
    doc = Document()
    set_page_margins(doc)

    add_header_table(doc, data)
    add_title(doc, f"Phân công nhiệm vụ cán bộ đoàn {data['loai_tu_van']}")
    add_tgd(doc)
    add_can_cu(doc, data)
    add_quyet_dinh_label(doc)

    p1 = build_dieu1_header(doc, "Phân công nhiệm vụ cán bộ", data["loai_tu_van"], data)
    fmt_body(p1.add_run(" gồm các cán bộ theo danh sách đính kèm, cụ thể:"))

    add_dieu_2_3_4(doc, data["ten_tt"], data["loai_tu_van"])
    add_footer_table(doc)

    # Trang 2 – bảng danh sách
    doc.add_page_break()
    _add_danh_sach_page(doc, data)
    return doc


def _add_danh_sach_page(doc, data):
    from builders.common import fmt_body as fb, para_spacing_body as psb

    p_ds = doc.add_paragraph()
    p_ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    psb(p_ds, align=WD_ALIGN_PARAGRAPH.CENTER)
    fb(p_ds.add_run("DANH SÁCH PHÂN CÔNG NHIỆM VỤ CÁN BỘ ĐOÀN TƯ VẤN"), bold=True)

    p_kt = doc.add_paragraph()
    psb(p_kt, align=WD_ALIGN_PARAGRAPH.CENTER)
    fb(p_kt.add_run(
        f"Kèm theo Quyết định số: {data['so_qd_display']} {data['ngay_qd_display']}"
    ))

    doc.add_paragraph()   # dòng trống

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, w in enumerate(COL_WIDTHS):
        table.columns[i].width = w
        table.rows[0].cells[i].width = w

    for i, h in enumerate(["STT", "Họ và tên", "Vị trí công việc"]):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        fb(p.add_run(h), bold=True)
        set_cell_vertical_align_center(table.rows[0].cells[i])

    for idx, m in enumerate(data["members"], 1):
        row = table.add_row()
        for i, w in enumerate(COL_WIDTHS):
            row.cells[i].width = w

        p_stt = row.cells[0].paragraphs[0]
        p_stt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p_stt, align=WD_ALIGN_PARAGRAPH.CENTER)
        fb(p_stt.add_run(str(idx)))
        set_cell_vertical_align_center(row.cells[0])

        p_ten = row.cells[1].paragraphs[0]
        para_spacing_table(p_ten)
        fb(p_ten.add_run(f"{m['trinh_do']}. {m['ho_ten']}"))
        set_cell_vertical_align_center(row.cells[1])

        p_cv = row.cells[2].paragraphs[0]
        para_spacing_table(p_cv)
        fb(p_cv.add_run(m["chuc_vu"]))
        set_cell_vertical_align_center(row.cells[2])
