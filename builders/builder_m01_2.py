from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from builders.common import *

def build(data):
    doc = Document()
    set_page_margins(doc)
    
    add_header_table(doc, data)
    add_title(doc, f"Thành lập đoàn {data['loai_tu_van']}")
    add_tgd(doc)
    add_can_cu(doc, data)
    add_quyet_dinh_label(doc)
    
    # Điều 1
    p1 = build_dieu1_header(doc, "Thành lập", data['loai_tu_van'], data)
    r_tail = p1.add_run(" gồm các cán bộ có tên trong danh sách đính kèm Quyết định này.")
    fmt_body(r_tail)
    
    add_dieu_2_3_4(doc, data["ten_tt"], data["loai_tu_van"])
    add_footer_table(doc)
    
    # Trang 2: Danh sách đính kèm
    doc.add_page_break()
    p_ds = doc.add_paragraph()
    p_ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing_body(p_ds, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_ds.add_run("DANH SÁCH ĐOÀN TƯ VẤN"), bold=True)
    
    p_kt = doc.add_paragraph()
    p_kt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing_body(p_kt, align=WD_ALIGN_PARAGRAPH.CENTER)
    fmt_body(p_kt.add_run(f"(Kèm theo Quyết định số: {data['so_qd_display']} ngày {data['ngay_qd_display']})"))
    
    # 1 dòng trống tạo khoảng cách trước bảng
    p_blank = doc.add_paragraph()
    para_spacing_body(p_blank)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    COL_WIDTHS = [Mm(15), Mm(85), Mm(60)]
    
    # Set chiều rộng header row
    hdr_cells = table.rows[0].cells
    col_headers = ["STT", "Họ và tên", "Vị trí công việc"]
    for i, width in enumerate(COL_WIDTHS):
        table.columns[i].width = width
        hdr_cells[i].width = width
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt_body(p.add_run(col_headers[i]), bold=True)
        set_cell_vertical_align_center(hdr_cells[i])

    # --- Data rows ---
    for idx, m in enumerate(data["members"], 1):
        row = table.add_row()
        
        # Set chiều rộng lại cho từng ô vừa thêm
        for i, width in enumerate(COL_WIDTHS):
            row.cells[i].width = width

        # Ô STT – căn giữa ngang + vertical center
        p_stt = row.cells[0].paragraphs[0]
        p_stt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing_table(p_stt, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt_body(p_stt.add_run(str(idx)))
        set_cell_vertical_align_center(row.cells[0])

        # Ô họ tên – căn trái + vertical center
        p_ten = row.cells[1].paragraphs[0]
        para_spacing_table(p_ten) # Mặc định giãn 3pt/3pt/15pt
        fmt_body(p_ten.add_run(f"{m['trinh_do']}. {m['ho_ten']}"))
        set_cell_vertical_align_center(row.cells[1])

        # Ô chức vụ – căn trái + vertical center
        p_cv = row.cells[2].paragraphs[0]
        para_spacing_table(p_cv)
        fmt_body(p_cv.add_run(m["chuc_vu"]))
        set_cell_vertical_align_center(row.cells[2])

    return doc
